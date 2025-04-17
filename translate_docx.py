from docx import Document
from tqdm import tqdm 
import time
from functools import partial
import argostranslate.package
import argostranslate.translate
import random 

from utils import preserve_special_elements

#imported by Mamun
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed


logging.basicConfig(format='%(levelname)s:%(message)s', level=logging.DEBUG)

def multiple_language_translation_setup(source_lang: str, target_lang: str):
    
    logging.info(f"translation package installetion checking: {source_lang} to {target_lang}")
    argostranslate.package.update_package_index()
    available_packages = argostranslate.package.get_available_packages()

    try:
        package_to_install = next(
        filter(
            lambda x: x.from_code == source_lang and x.to_code == target_lang, available_packages
        )
        )
        argostranslate.package.install_from_path(package_to_install.download())
    except Exception as e:
        logging.error(f"Found an Exception: {e} for {source_lang} to {target_lang} translation package")
        raise ValueError(f"translation packege or model is not available for {source_lang} to {target_lang}")


def translate(source_lang: str, target_lang: str, verbose=False) -> str:
    """
    Translate text using argostranslate.

    Args:
        text (str): The text to translate
        source_lang (str): The source language of the text
        target_lang (str): The target language to translate to

    Returns:
        str: The translated text

    Note:
        - Uses MODEL_NAME and MODEL_URL from environment variables
        - Logs translations to 'translations.txt'
    """
    # TODO (ex.2) : You can modify this function except line 41-42
    # Download and install Argos Translate package
    # argostranslate.package.update_package_index()
    # available_packages = argostranslate.package.get_available_packages()
    # package_to_install = next(
    #     filter(
    #         lambda x: x.from_code == source_lang and x.to_code == target_lang, available_packages
    #     )
    # )
    # argostranslate.package.install_from_path(package_to_install.download())

    multiple_language_translation_setup(source_lang, target_lang)
    def translation_func(text: str) -> str:
        ############ DO NOT CHANGE THIS BLOCK OF CODE ##############################
        # Simulate HTTP request latency with a random delay between 0.1 and 0.4 seconds
        delay = random.uniform(0.1, 0.4)
        time.sleep(delay)
        #############################################################################
        
        # Translate
        response = argostranslate.translate.translate(text, source_lang, target_lang)
        
        if verbose : 
            with open('translations.txt', 'a') as f:
                f.write(text + '=>'+ response + '\n')
        return response
    return translation_func


def translate_runs_in_paragraph(paragraph: any, translation_func: callable, mode: str = "naive") -> None:
    """
    TODO (ex.1) : finish this doc
    """
    # First preserve special elements with their positions
    special_elements = preserve_special_elements(paragraph)

    if mode == "naive":
        for run in paragraph.runs:
            if run.text.strip() and len(run.text) < 2000: # dodgy way to avoid translating long tables/graphs 
                translated_text = translation_func(run.text)
                run.text = translated_text.strip()+" "

    # Restore special elements in their original positions
    for original_pos, element in special_elements:
        # Insert the special element back at its relative position
        paragraph._p.insert(original_pos, element)

                

def translate_docx(input_path, output_path, source_lang, target_lang):
    """
    TODO (ex.1) : finish this doc
    I could use parallan processing in this function to reduce execution time.
    """
    # TODO (ex.2) : You can modify this whole function to make the process faster
    # Open the document
    start = time.time()

    translation_func  = translate(source_lang, target_lang)#partial(translate, source_lang=source_lang, target_lang=target_lang)

    #print(f"Opening document: {input_path}")
    logging.info(f"Opening document: {input_path}")
    doc = Document(input_path)
    
    #print(f"Translation in progress...")
    logging.info("Translating paragraphs...")
    paragraphs_to_translate = []
    
    for paragraph in tqdm(doc.paragraphs):
        #translate_runs_in_paragraph(paragraph, translation_func)
        paragraphs_to_translate.append(paragraph)

    for table in tqdm(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    #translate_runs_in_paragraph(paragraph, translation_func)
                    paragraphs_to_translate.append(paragraph)


    for section in tqdm(doc.sections):
        for paragraph in section.header.paragraphs:
            #translate_runs_in_paragraph(paragraph, translation_func)
            paragraphs_to_translate.append(paragraph)
        for paragraph in  section.footer.paragraphs:
            #translate_runs_in_paragraph(paragraph, translation_func)
            paragraphs_to_translate.append(paragraph)
        for table in  section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in tqdm(cell.paragraphs):
                        #translate_runs_in_paragraph(paragraph, translation_func)
                        paragraphs_to_translate.append(paragraph)
    

    with ThreadPoolExecutor() as executor:
        futures = []
        for paragraph in paragraphs_to_translate:
            futures.append(executor.submit(translate_runs_in_paragraph, paragraph, translation_func))

        for future in as_completed(futures):
            future.result()  

    #print(f"Translation is done !")
    logging.info("Translation is done !")
    #print(f"Saving translated document to: {output_path}")
    logging.info(f"Saving translated document to: {output_path}")
    doc.save(output_path)
    end = time.time()
    #print(f"Total time taken: {end - start} seconds.")
    logging.info(f"Document translated successfully in {end - start:.2f} seconds.")


def main() -> None:
    """
    Main entry point for document translation.
    Translates a sample document from English to French Canadian.
    """

    input_path = "sample_docs/brochure.docx" # Any BMO doc would work here 
    output_path = input_path.replace(".docx", "_fr.docx")
    translate_docx(input_path, output_path, source_lang="en", target_lang="fr")
    print(f"Document translated and saved to '{output_path}'.")

if __name__=="__main__":

    main()
 